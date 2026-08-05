# -*- coding: utf-8 -*-
"""docx 排版与引用标记复用函数。
用法: from docx_helper import new_doc, add_heading, add_field, add_para, add_mixed
引用纪律:
  - 原文引用 style='q' → 楷体、蓝色 RGB(0x1F,0x4E,0x79)
  - 改写概括 style='g' → 仿宋、灰色 RGB(0x2E,0x2E,0x2E)
  - 加粗 style='b' / 普通 style='n'
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

Q_COLOR = (0x1F, 0x4E, 0x79)  # 原文引用：蓝
G_COLOR = (0x2E, 0x2E, 0x2E)  # 改写概括：灰


def set_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text='', size=12, bold=False, align=None, font='宋体',
             space_after=6, indent=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.3
    if indent:
        pf.first_line_indent = Pt(indent)
    if text:
        r = p.add_run(text)
        set_font(r, name=font, size=size, bold=bold, color=color)
    return p


def add_mixed(doc, segments, size=12, indent=None, space_after=6):
    """segments: list of (text, style); style ∈ {'q'|原文引用, 'g'|改写, 'b'|加粗, 'n'|普通}"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.3
    if indent:
        pf.first_line_indent = Pt(indent)
    for text, style in segments:
        r = p.add_run(text)
        if style == 'q':
            set_font(r, name='楷体', size=size, color=Q_COLOR)
        elif style == 'g':
            set_font(r, name='仿宋', size=size, color=G_COLOR)
        elif style == 'b':
            set_font(r, size=size, bold=True)
        else:
            set_font(r, size=size)
    return p


def add_heading(doc, text, size=16):
    return add_para(doc, text, size=size, bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER, font='黑体', space_after=10)


def add_field(doc, label, value, bold_label=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.3
    r1 = p.add_run(label)
    set_font(r1, size=12, bold=bold_label)
    r2 = p.add_run(value)
    set_font(r2, size=12)
    return p


def new_doc():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.8)
    return doc
